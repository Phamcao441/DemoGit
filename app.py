import os, json, threading, time, concurrent.futures
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_file
from pdf2image import convert_from_path
import pytesseract, cv2, numpy as np
from PIL import Image
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv

# Import từ models.py để đảm bảo kiến trúc DB 15 trường Dublin Core
from models import db, OCRRecord, TesseractResult, DublinCoreMetadata

# 1. CẤU HÌNH HỆ THỐNG
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = Flask(__name__)
app.config.update(
    SQLALCHEMY_DATABASE_URI="sqlite:///ocr.db",
    SQLALCHEMY_TRACK_MODIFICATIONS=False,
    UPLOAD_FOLDER='uploads',
    OUTPUT_FOLDER='outputs'
)
db.init_app(app)

# Cấu hình đường dẫn trên VPS Hải Phòng
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
POPPLER_PATH = r'C:\poppler\Library\bin'

for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
    os.makedirs(folder, exist_ok=True)

# Bộ nhớ tạm để Polling 30s - Giải pháp "bất tử" qua Cloudflare
tasks = {}

# 2. XỬ LÝ ẢNH TẦNG SÂU (IMAGE PREPROCESSING)
def preprocess_image(img):
    """Sử dụng thuật toán FastNLMeans và CLAHE để cứu dấu thanh tiếng Việt ở DPI 500"""
    img_np = np.array(img)
    gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
    
    # Khử nhiễu hạt để Tesseract không đọc nhầm dấu nặng/dấu chấm
    denoised = cv2.fastNlMeansDenoising(gray, h=10)
    
    # Cân bằng độ sáng cục bộ giúp hiện rõ các nét chữ bị mờ
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(denoised)
    
    # Ngưỡng thích nghi blockSize=51 giúp giữ nét chữ to rõ ở độ phân giải cao
    binary = cv2.adaptiveThreshold(
        enhanced, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY, 51, 11
    )
    return Image.fromarray(binary)

def ocr_page(i, img):
    processed = preprocess_image(img)
    # Chế độ psm 6 phù hợp cho văn bản sách của bác Phan Văn Trường
    text = pytesseract.image_to_string(processed, lang='vie', config='--psm 6 --oem 3')
    return i, text.strip()

# 3. AI PHỤC CHẾ KỶ LUẬT THÉP (ANTI-HALLUCINATION)
def ai_reconstruct_batch(text):
    """Ép AI phục chế trung thực, cấm thay đổi từ ngữ gốc"""
    if not text.strip(): return ""
    try:
        res = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system", 
                    "content": (
                        "BẠN LÀ MỘT HỆ THỐNG PHỤC CHẾ VĂN BẢN (TEXT RESTORER). "
                        "QUY TẮC CẤM: 1. TUYỆT ĐỐI KHÔNG THAY ĐỔI TỪ VỰNG CỦA TÁC GIẢ (Ví dụ: Giữ nguyên 'đồng tính', KHÔNG đổi thành 'sống thử'). "
                        "2. KHÔNG tóm tắt, không viết lại câu cho hay hơn. 3. CHỈ PHỤC CHẾ KÝ TỰ: Nhiệm vụ của bạn là 'khâu vá' các nét chữ bị vỡ (rn -> m, d -> đ, ol -> ơi), không phải là biên tập viên văn học.. "
                        "4. Trả về 100% nội dung gốc dưới dạng văn bản thuần túy."
                    )
                },
                {"role": "user", "content": f"Hãy phục chế đoạn OCR thô sau:\n\n{text}"}
            ],
            temperature=0 # Ép AI chạy theo logic cứng, không sáng tạo
        )
        return res.choices[0].message.content
    except: return text

# 4. BACKGROUND WORKER
def background_worker(task_id, path, filename):
    try:
        tasks[task_id] = {"status": "Đang rã PDF siêu nét...", "progress": 0, "total": 0, "partial": "", "done": False}
        
        # Chuyển đổi PDF sang ảnh chất lượng cao 500 DPI
        images = convert_from_path(path, dpi=500, poppler_path=POPPLER_PATH)
        total = len(images)
        tasks[task_id]["total"] = total
        raw_pages = [""] * total

        # OCR đa luồng (Dùng 2-3 luồng để an toàn cho RAM VPS)
        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
            futures = [executor.submit(ocr_page, i, img) for i, img in enumerate(images)]
            for future in concurrent.futures.as_completed(futures):
                idx, text = future.result()
                raw_pages[idx] = text
                tasks[task_id]["progress"] = idx + 1
                tasks[task_id]["status"] = f"Xeon đang quét: {idx+1}/{total}"

        # AI Phục chế mẻ 5 trang để giữ ngữ cảnh
        tasks[task_id]["status"] = "AI đang phục chế trung thành..."
        corrected_chunks = []
        for i in range(0, total, 5):
            batch = raw_pages[i : i+5]
            ai_text = ai_reconstruct_batch("\n\n".join(batch))
            corrected_chunks.append(ai_text)
            tasks[task_id]["partial"] += ai_text + "\n\n"

        full_text = "\n\n".join(corrected_chunks)

        # Trích xuất 15 trường Dublin Core cho Metadata
        res_meta = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "JSON 15 trường Dublin Core: title, creator, subject, description, publisher, contributor, date, type, format, identifier, source, language, relation, coverage, rights"},
                {"role": "user", "content": full_text[:6000]}
            ],
            response_format={"type": "json_object"}
        )
        meta = json.loads(res_meta.choices[0].message.content)

        tasks[task_id].update({
            "status": "SUCCESS", "done": True, "metadata": meta,
            "full_text": full_text, "raw_full": "\n\n".join(raw_pages),
            "filename": filename, "path": path
        })
    except Exception as e:
        tasks[task_id] = {"status": "FAILED", "error": str(e), "done": True}

# 5. ROUTES
@app.route('/')
def index(): return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    file = request.files['file']
    task_id = datetime.now().strftime('%H%M%S%f')
    path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(path)
    threading.Thread(target=background_worker, args=(task_id, path, file.filename)).start()
    return jsonify({"task_id": task_id})

@app.route('/status/<task_id>')
def status(task_id): 
    # API Polling 30s cho trình duyệt
    return jsonify(tasks.get(task_id, {"status": "NOT_FOUND"}))

@app.route('/finalize', methods=['POST'])
def finalize():
    """Ghi dữ liệu đã hiệu chỉnh vào DB và xuất Word"""
    data = request.json
    task_id = data.get('task_id')
    user_meta = data.get('metadata')
    task = tasks.get(task_id)
    
    try:
        with app.app_context():
            new_rec = OCRRecord(filename=task['filename'], input_path=task['path'], content=task['full_text'])
            db.session.add(new_rec)
            db.session.flush()
            
            db.session.add(TesseractResult(ocr_id=new_rec.id, raw_content=task['raw_full']))
            db.session.add(DublinCoreMetadata(ocr_id=new_rec.id, **user_meta))
            db.session.commit()

            doc_path = os.path.join(app.config['OUTPUT_FOLDER'], f"{task_id}.docx")
            doc = Document()
            doc.add_heading(user_meta.get('title', 'KẾT QUẢ SỐ HÓA'), 0)
            doc.add_paragraph(task['full_text'])
            doc.save(doc_path)
            
        return jsonify({"success": True, "download": f"/download/{task_id}.docx"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/download/<f>')
def download(f): return send_file(os.path.join(app.config['OUTPUT_FOLDER'], f), as_attachment=True)

if __name__ == '__main__':
    with app.app_context():
        db.create_all() # Tự động tạo cấu hình 15 trường Metadata
    # Mở host 0.0.0.0 để đệ truy cập từ xa qua IP VPS
    app.run(host='0.0.0.0', port=8080, debug=True)